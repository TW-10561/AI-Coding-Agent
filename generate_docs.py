#!/usr/bin/env python3
"""
Generate DOCX documentation for Thirdwave AI Coding Platform
Creates 60_xx manual documents with Japanese-English bilingual content
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

def add_heading_style(doc, text, level=1):
    """Add styled heading"""
    heading = doc.add_heading(text, level=level)
    return heading

def add_table_style(doc, rows, cols, data):
    """Add table with data"""
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Light Grid Accent 1'
    for i, row_data in enumerate(data):
        for j, cell_data in enumerate(row_data):
            table.rows[i].cells[j].text = str(cell_data)
    return table

def set_cell_background(cell, fill_color):
    """Set table cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill_color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def create_60_01_admin_manual():
    """60_01: Administrator Manual (管理者マニュアル)"""
    doc = Document()
    
    # Title
    title = doc.add_paragraph()
    title_run = title.add_run('Thirdwave AI Coding Platform\n管理者マニュアル')
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f'Administrator Manual / バージョン 0.1.0\n生成日: {datetime.now().strftime("%Y年%m月%d日")}')
    
    # Table of Contents
    add_heading_style(doc, '目次 (Table of Contents)', 1)
    toc_items = [
        '1. システム概要 (System Overview)',
        '2. インストール手順 (Installation)',
        '3. デプロイメント (Deployment)',
        '4. 設定管理 (Configuration Management)',
        '5. 監視とメンテナンス (Monitoring & Maintenance)',
        '6. バックアップと復旧 (Backup & Recovery)',
        '7. トラブルシューティング (Troubleshooting)',
        '8. セキュリティポリシー (Security Policies)',
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Bullet')
    
    # Section 1
    add_heading_style(doc, '1. システム概要 (System Overview)', 1)
    doc.add_paragraph(
        'Thirdwave AI Coding Platform は、ローカルに展開可能なAIコーディング支援システムです。'
        'vLLMなどのローカルGPU推論、複数のクラウドAIプロバイダーとの統合、'
        'Human-in-the-Loop（HITL）ガバナンス、エンタープライズグレードの監査ログを提供します。'
    )
    doc.add_paragraph(
        'The Thirdwave AI Coding Platform is a self-hosted AI coding assistant with local GPU inference support, '
        'cloud provider integration, HITL governance, and enterprise-grade audit logging.'
    )
    
    add_heading_style(doc, '1.1 アーキテクチャ (Architecture)', 2)
    arch_table_data = [
        ['コンポーネント', 'ポート', '説明'],
        ['Platform API', '3100', 'Honoベースのバックエンドサーバー'],
        ['OpenCode Engine', '4096', 'AIセッション管理とツール実行'],
        ['nginx Proxy', '80/443', 'リバースプロキシとレート制限'],
        ['vLLM', '8000', 'ローカルGPU推論'],
        ['AI Gateway', '9080', 'クラウドプロバイダルーティング'],
    ]
    add_table_style(doc, 6, 3, arch_table_data)
    
    # Section 2
    add_heading_style(doc, '2. インストール手順 (Installation)', 1)
    
    add_heading_style(doc, '2.1 前提条件 (Requirements)', 2)
    reqs = [
        'Bun >= 1.3.0',
        'OpenCode >= 1.2.17',
        'GPU サーバー: vLLM >= 0.5.0 または Ollama >= 0.5.0',
        'nginx >= 1.18（オプション、ポート80アクセス用）',
        'systemd（サービス管理用）',
    ]
    for req in reqs:
        doc.add_paragraph(req, style='List Bullet')
    
    add_heading_style(doc, '2.2 開発環境セットアップ (Development Setup)', 2)
    doc.add_paragraph('bash platform/scripts/dev-setup.sh')
    doc.add_paragraph('このスクリプトはBun、依存関係をインストールします。')
    
    add_heading_style(doc, '2.3 本番環境デプロイ (Production Deployment)', 2)
    doc.add_paragraph('1. Git リポジトリをクローン')
    doc.add_paragraph('2. .env ファイルを設定（VLLM_GATEWAY_URL, APIキーなど）', style='List Number')
    doc.add_paragraph('3. デプロイスクリプト実行: sudo bash platform/deploy/deploy.sh', style='List Number')
    doc.add_paragraph('4. ステータス確認: sudo bash platform/deploy/deploy.sh --status', style='List Number')
    
    # Section 3
    add_heading_style(doc, '3. デプロイメント (Deployment)', 1)
    
    add_heading_style(doc, '3.1 Docker デプロイ', 2)
    doc.add_paragraph('docker compose -f platform/docker/docker-compose.yml up --build')
    
    add_heading_style(doc, '3.2 nginx 設定 (nginx Configuration)', 2)
    nginx_table_data = [
        ['機能', '制限値'],
        ['一般API', '30 req/s（バースト50）'],
        ['チャット', '5 req/s（バースト10）'],
        ['接続制限', 'IP当たり20同時接続'],
        ['読み取りタイムアウト', '/health: 5-10s, /api/chat: 300s, SSE: 600s'],
    ]
    add_table_style(doc, 5, 2, nginx_table_data)
    
    add_heading_style(doc, '3.3 systemd サービス (systemd Service)', 2)
    doc.add_paragraph('systemd service は以下の設定で自動起動します:')
    doc.add_paragraph('• ユーザー: nvidia')
    doc.add_paragraph('• 再起動ポリシー: on-failure (5秒待機)')
    doc.add_paragraph('• タイムアウト: 30秒')
    
    # Section 4
    add_heading_style(doc, '4. 設定管理 (Configuration Management)', 1)
    
    add_heading_style(doc, '4.1 環境変数 (Environment Variables)', 2)
    config_table_data = [
        ['変数', 'デフォルト', '説明'],
        ['PORT', '3100', 'Platform API ポート'],
        ['VLLM_GATEWAY_URL', '-', 'AIゲートウェイエンドポイント'],
        ['VLLM_MODEL_ID', 'plezan/MiniMax...', 'デフォルトモデル'],
        ['PLATFORM_API_KEY', '-', '認証トークン（設定時のみ必須）'],
        ['LOG_LEVEL', 'info', 'ログレベル（debug/info/warn/error）'],
    ]
    add_table_style(doc, 6, 3, config_table_data)
    
    # Section 5
    add_heading_style(doc, '5. 監視とメンテナンス (Monitoring & Maintenance)', 1)
    
    add_heading_style(doc, '5.1 ヘルスチェック (Health Checks)', 2)
    doc.add_paragraph('curl http://localhost:3100/health/ready')
    doc.add_paragraph('このエンドポイントはDockerやロードバランサーの監視に使用します。')
    
    add_heading_style(doc, '5.2 ログ管理 (Log Management)', 2)
    doc.add_paragraph('• nginx ログ: /var/log/nginx/thirdwave-{access,error}.log')
    doc.add_paragraph('• Platform ログ: stdout（systemd journalctl で確認可能）')
    doc.add_paragraph('• 監査ログ: SQLite (.platform/audit.db)')
    
    # Section 6
    add_heading_style(doc, '6. バックアップと復旧 (Backup & Recovery)', 1)
    doc.add_paragraph(
        'SQLite WAL モードの.platform データベース、セッション状態、監査ログをバックアップしてください。'
    )
    doc.add_paragraph('毎日の自動バックアップを推奨: crontab -e で設定')
    
    # Section 7
    add_heading_style(doc, '7. トラブルシューティング (Troubleshooting)', 1)
    
    add_heading_style(doc, '7.1 ポート競合 (Port Conflict)', 2)
    doc.add_paragraph('AUTO_PORT=true を設定すると、自動的に空いているポートを探します。')
    
    add_heading_style(doc, '7.2 vLLM 接続不可 (vLLM Connection Failed)', 2)
    doc.add_paragraph('1. vLLM サービスが起動しているか確認: curl http://GPU_SERVER:8000/v1/models')
    doc.add_paragraph('2. ファイアウォール設定を確認', style='List Number')
    doc.add_paragraph('3. VLLM_GATEWAY_URL が正しく設定されているか確認', style='List Number')
    
    # Section 8
    add_heading_style(doc, '8. セキュリティポリシー (Security Policies)', 1)
    
    add_heading_style(doc, '8.1 認証 (Authentication)', 2)
    doc.add_paragraph('PLATFORM_API_KEY を設定することで、すべての /api/* エンドポイントに認証が必須になります。')
    doc.add_paragraph('クライアントは Authorization: Bearer <TOKEN> ヘッダーを含める必要があります。')
    
    add_heading_style(doc, '8.2 HITL ガバナンス (HITL Governance)', 2)
    doc.add_paragraph('以下のアクション は自動承認されず、ユーザー確認が必要です（ロールに応じて）:')
    doc.add_paragraph('• 破壊的なシェルコマンド (rm -rf, DROP TABLE など)', style='List Bullet')
    doc.add_paragraph('• 機密ファイルアクセス (.env, SSH キーなど)', style='List Bullet')
    doc.add_paragraph('• 外部ネットワークアクセス', style='List Bullet')
    
    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph(f'© 2026 Thirdwave AI Platform - 管理者マニュアル')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.save('/home/nvidia/AI_Coding_Agent/Kadavuley/AI-Coding-Agent/platform/docs/manual/60_01_管理者マニュアル.docx')
    print('✓ 60_01_管理者マニュアル.docx 生成完了')

def create_60_02_user_manual():
    """60_02: User / Operation Manual (操作マニュアル)"""
    doc = Document()
    
    # Title
    title = doc.add_paragraph()
    title_run = title.add_run('Thirdwave AI Coding Platform\n操作マニュアル')
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f'User / Operation Manual / バージョン 0.1.0\n生成日: {datetime.now().strftime("%Y年%m月%d日")}')
    
    # TOC
    add_heading_style(doc, '目次 (Table of Contents)', 1)
    toc = ['1. 基本操作 (Getting Started)', '2. チャット操作 (Chat Operations)', 
           '3. ファイル管理 (File Management)', '4. モデル選択 (Model Selection)',
           '5. スキル活用 (Skills & Agents)', '6. セッション管理 (Session Management)',
           '7. VS Code 拡張機能 (VS Code Extension)']
    for item in toc:
        doc.add_paragraph(item, style='List Bullet')
    
    # Section 1
    add_heading_style(doc, '1. 基本操作 (Getting Started)', 1)
    
    add_heading_style(doc, '1.1 起動方法 (Launch)', 2)
    doc.add_paragraph('TUI（ターミナルUI）で起動:')
    doc.add_paragraph('cd platform && bun run launch')
    doc.add_paragraph()
    doc.add_paragraph('ヘッドレス（バックエンドのみ）で起動:')
    doc.add_paragraph('bun run platform/scripts/start-all.ts')
    
    add_heading_style(doc, '1.2 初期設定 (Initial Setup)', 2)
    doc.add_paragraph('1. ステータスコマンドを実行: /status', style='List Number')
    doc.add_paragraph('2. 利用可能なモデルを確認: /registry', style='List Number')
    doc.add_paragraph('3. エージェントを選択: /build（デフォルト）またはその他', style='List Number')
    
    # Section 2
    add_heading_style(doc, '2. チャット操作 (Chat Operations)', 1)
    
    add_heading_style(doc, '2.1 メッセージ送信 (Sending Messages)', 2)
    doc.add_paragraph(
        'TUI: テキスト入力 → Enter で送信\n'
        'コピペ対応: 複数行の貼り付けは自動で集約されます\n'
        'VS Code: @thirdwave コマンドパレットか Chat パネルで入力'
    )
    
    add_heading_style(doc, '2.2 コマンド一覧 (Commands List)', 2)
    cmd_table = [
        ['コマンド', '機能', '例'],
        ['/explain', 'コードを説明', '/explain <code>'],
        ['/fix', 'バグを修正', '/fix <error>'],
        ['/test', 'テストを生成', '/test <function>'],
        ['/review', 'コード審査', '/review <file>'],
        ['/status', 'システムステータス', '/status'],
        ['/models', 'モデル一覧', '/models'],
        ['/skills', 'スキル一覧', '/skills'],
        ['/new', '新規セッション', '/new'],
        ['/help', 'ヘルプ表示', '/help'],
    ]
    add_table_style(doc, len(cmd_table), 3, cmd_table)
    
    # Section 3
    add_heading_style(doc, '3. ファイル管理 (File Management)', 1)
    
    add_heading_style(doc, '3.1 ファイル操作 (File Operations)', 2)
    doc.add_paragraph('読込: /files で現在の構造を表示')
    doc.add_paragraph('プロジェクト情報: /project で詳細を確認')
    doc.add_paragraph('VCS 情報: /vcs で Git ブランチ等を確認')
    
    add_heading_style(doc, '3.2 AI とのファイル共有 (Sharing with AI)', 2)
    doc.add_paragraph(
        'VS Code 拡張機能では、選択したコードを自動的に AI に送信できます。\n'
        'チャット内で /explain, /fix, /review コマンドを使用してください。'
    )
    
    # Section 4
    add_heading_style(doc, '4. モデル選択 (Model Selection)', 1)
    
    add_heading_style(doc, '4.1 利用可能なモデル (Available Models)', 2)
    doc.add_paragraph(
        'ローカル GPU: vLLM から自動検出 (Qwen3-Coder, MiniMax など)\n'
        'クラウド: OpenAI, Anthropic, Google AI, Mistral など'
    )
    
    add_heading_style(doc, '4.2 モデル変更 (Changing Model)', 2)
    doc.add_paragraph('TUI: /models コマンドで表示 → 選択')
    doc.add_paragraph('VS Code: プルダウンからモデルを選択')
    doc.add_paragraph('トークン制限: リフレッシュボタンで最新データを取得')
    
    # Section 5
    add_heading_style(doc, '5. スキル活用 (Skills & Agents)', 1)
    
    add_heading_style(doc, '5.1 スキルシステム (Skills System)', 2)
    doc.add_paragraph(
        '48+ の専門スキルが組み込まれています。カテゴリ:\n'
        '• 開発: API設計, アーキテクチャパターン, テスト駆動開発\n'
        '• Azure: AI/Deploy/監視/最適化\n'
        '• DevOps: CI/CD, デプロイメント自動化\n'
        '• フロントエンド: React, UI/UX設計'
    )
    
    add_heading_style(doc, '5.2 エージェント選択 (Agent Selection)', 2)
    agent_table = [
        ['エージェント', '用途'],
        ['build', 'コード作成・デバッグ'],
        ['test', 'テスト作成'],
        ['review', 'コード審査'],
        ['explore', '調査・ドキュメント'],
        ['general', '汎用タスク'],
    ]
    add_table_style(doc, len(agent_table), 2, agent_table)
    
    # Section 6
    add_heading_style(doc, '6. セッション管理 (Session Management)', 1)
    
    add_heading_style(doc, '6.1 セッション操作 (Session Operations)', 2)
    doc.add_paragraph('/new - 新しいセッション作成')
    doc.add_paragraph('/sessions - 履歴表示', style='List Bullet')
    doc.add_paragraph('/switch - セッション切り替え', style='List Bullet')
    doc.add_paragraph('/delete - セッション削除', style='List Bullet')
    
    add_heading_style(doc, '6.2 セッション永続化 (Session Persistence)', 2)
    doc.add_paragraph('セッションは自動的に保存されます（SQLite）。')
    doc.add_paragraph('再起動後も履歴が保持されます。')
    
    # Section 7
    add_heading_style(doc, '7. VS Code 拡張機能 (VS Code Extension)', 1)
    
    add_heading_style(doc, '7.1 インストール (Installation)', 2)
    doc.add_paragraph('1. VS Code の拡張機能パネルで "Thirdwave AI" を検索')
    doc.add_paragraph('2. インストール & 有効化', style='List Number')
    doc.add_paragraph('3. チャットパレット (@thirdwave) で開始', style='List Number')
    
    add_heading_style(doc, '7.2 主な機能 (Features)', 2)
    doc.add_paragraph('• チャットパネル: @thirdwave で AI と会話')
    doc.add_paragraph('• スラッシュコマンド: /explain, /fix, /test, /review, /models')
    doc.add_paragraph('• モデルピッカー: ドロップダウンで数十の AI モデルを選択')
    doc.add_paragraph('• スキル統合: 48+ の domain skills が自動で活用される')
    
    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph(f'© 2026 Thirdwave AI Platform - 操作マニュアル')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.save('/home/nvidia/AI_Coding_Agent/Kadavuley/AI-Coding-Agent/platform/docs/manual/60_02_操作マニュアル.docx')
    print('✓ 60_02_操作マニュアル.docx 生成完了')

def create_60_03_readme():
    """60_03: Readme / Setup Overview (Readme)"""
    doc = Document()
    
    # Title
    title = doc.add_paragraph()
    title_run = title.add_run('Thirdwave AI Coding Platform\nREADME / セットアップガイド')
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f'Readme & Setup Overview / バージョン 0.1.0\n生成日: {datetime.now().strftime("%Y年%m月%d日")}')
    
    # Overview
    add_heading_style(doc, 'プロジェクト概要 (Project Overview)', 1)
    doc.add_paragraph(
        'Thirdwave AI Coding Platform は、エンタープライズグレードの自己ホスト型 AI コーディング支援プラットフォームです。'
        'ローカル GPU 推論（vLLM）とクラウド AI プロバイダー（OpenAI, Anthropic など）の両方に対応しており、'
        'Human-in-the-Loop ガバナンス、エンタープライズ監査ログ、マルチエージェント オーケストレーション、'
        '継続的な改善ループなどの高度な機能を備えています。'
    )
    
    # Quick Start
    add_heading_style(doc, 'クイックスタート (Quick Start)', 1)
    
    add_heading_style(doc, 'ステップ 1: リポジトリをクローン', 2)
    doc.add_paragraph('git clone <repository_url>')
    doc.add_paragraph('cd AI-Coding-Agent/platform')
    
    add_heading_style(doc, 'ステップ 2: .env ファイルを作成', 2)
    doc.add_paragraph('cp .env.example .env')
    doc.add_paragraph('# 以下の環境変数を設定:')
    doc.add_paragraph('# VLLM_GATEWAY_URL=http://your-gpu-server:8000')
    doc.add_paragraph('# VLLM_MODEL_ID=Qwen3-Coder-30B')
    
    add_heading_style(doc, 'ステップ 3: 起動', 2)
    doc.add_paragraph('bun run launch')
    
    # System Architecture
    add_heading_style(doc, 'システムアーキテクチャ (System Architecture)', 1)
    
    architecture = """
    ┌─────────────────────────────────────┐
    │     VS Code Extension / TUI         │
    │  (ユーザーインターフェース)         │
    └────────────────┬────────────────────┘
                     │
    ┌────────────────▼────────────────────┐
    │     Platform API (Hono)             │ ← ポート 3100
    │   (Middleware, Routes, Services)    │
    └────────────────┬────────────────────┘
                     │
    ┌────────────────▼────────────────────┐
    │   OpenCode Engine                   │ ← ポート 4096
    │ (AI Session, Tool Executor, Loop)   │
    └────────────────┬────────────────────┘
              ┌──────┴──────┐
              │             │
    ┌─────────▼──┐  ┌──────▼────────┐
    │  vLLM      │  │ Cloud AI       │
    │  (Local)   │  │ (OpenAI, etc)  │
    └────────────┘  └────────────────┘
    """
    doc.add_paragraph(architecture)
    
    # Requirements
    add_heading_style(doc, '必要な環境 (Requirements)', 1)
    
    add_heading_style(doc, 'ソフトウェア (Software)', 2)
    req_table = [
        ['項目', '最小バージョン', '推奨'],
        ['Bun', '1.3.0', '1.3.10+'],
        ['OpenCode', '1.2.17', '1.2.17+'],
        ['Node/TypeScript', '18+/5.x', 'Bun で自動'],
        ['Python', '3.8+', '3.11+'],
        ['bash', '4.0+', '5.x+'],
    ]
    add_table_style(doc, len(req_table), 3, req_table)
    
    add_heading_style(doc, 'ハードウェア (Hardware)', 2)
    doc.add_paragraph('• CPU: 4+ cores 推奨')
    doc.add_paragraph('• RAM: 8GB 以上（16GB+ 推奨）')
    doc.add_paragraph('• GPU: CUDA/ROCm 対応（推奨）— CPU単独でも動作')
    doc.add_paragraph('• ストレージ: 20GB 以上の空きスペース')
    
    # Deployment Options
    add_heading_style(doc, 'デプロイメントオプション (Deployment Options)', 1)
    
    add_heading_style(doc, '1. 開発モード (Development)', 2)
    doc.add_paragraph('bun run launch')
    doc.add_paragraph('• OpenCode + Platform + TUI がすべて起動')
    doc.add_paragraph('• ローカル開発用')
    
    add_heading_style(doc, '2. ヘッドレスモード (Headless)', 2)
    doc.add_paragraph('bun run start:all')
    doc.add_paragraph('• Backend のみ起動（TUI なし）')
    doc.add_paragraph('• Docker / systemd 用')
    
    add_heading_style(doc, '3. Docker デプロイ (Docker)', 2)
    doc.add_paragraph('docker compose -f platform/docker/docker-compose.yml up')
    doc.add_paragraph('• コンテナ化されたデプロイメント')
    doc.add_paragraph('• 本番環境推奨')
    
    add_heading_style(doc, '4. systemd サービス (systemd Service)', 2)
    doc.add_paragraph('sudo bash platform/deploy/deploy.sh')
    doc.add_paragraph('• nginx + systemd service をセットアップ')
    doc.add_paragraph('• 自動再起動、persistent logging')
    
    # Configuration
    add_heading_style(doc, '設定 (Configuration)', 1)
    
    add_heading_style(doc, '主要な環境変数 (Key Environment Variables)', 2)
    config_table = [
        ['変数', 'デフォルト', '説明'],
        ['PORT', '3100', 'Backend ポート'],
        ['VLLM_GATEWAY_URL', '-', 'GPU 推論エンドポイント'],
        ['VLLM_MODEL_ID', 'plezan/MiniMax...', 'デフォルト AI モデル'],
        ['OPENAI_API_KEY', '-', 'OpenAI 統合（オプション）'],
        ['LOG_LEVEL', 'info', 'ログ詳細度'],
    ]
    add_table_style(doc, len(config_table), 3, config_table)
    
    # Features
    add_heading_style(doc, '主要機能 (Key Features)', 1)
    
    features_list = [
        '✓ ローカル & クラウド AI モデル統合',
        '✓ マルチエージェント オーケストレーション',
        '✓ Human-in-the-Loop ガバナンス & 自動リスク評価',
        '✓ エンタープライズ監査ログ',
        '✓ 48+ ドメイン スキル (API設計, テスト駆動開発など)',
        '✓ VS Code 拡張機能',
        '✓ RESTful API + SSE リアルタイム ストリーミング',
        '✓ SQLite ベースのセッション永続化',
        '✓ 並列実行 & 非同期タスクキュー',
        '✓ ロールベースアクセス制御 (RBAC)',
    ]
    for feature in features_list:
        doc.add_paragraph(feature, style='List Bullet')
    
    # Troubleshooting
    add_heading_style(doc, 'トラブルシューティング (Troubleshooting)', 1)
    
    add_heading_style(doc, '問題: "vLLM に接続できない"', 2)
    doc.add_paragraph('1. GPU サーバーが起動しているか確認: curl http://GPU:8000/v1/models')
    doc.add_paragraph('2. VLLM_GATEWAY_URL が正しいか確認', style='List Number')
    doc.add_paragraph('3. ファイアウォール設定を確認', style='List Number')
    
    add_heading_style(doc, '問題: "ポート 3100 が使用中"', 2)
    doc.add_paragraph('(自動) AUTO_PORT=true を設定すると別ポートを自動選択')
    doc.add_paragraph('(手動) PORT=3101 bun run launch', style='List Bullet')
    
    # Support & Documentation
    add_heading_style(doc, 'ドキュメント & サポート (Support)', 1)
    doc.add_paragraph('• 管理者マニュアル: 60_01_管理者マニュアル.docx')
    doc.add_paragraph('• 操作マニュアル: 60_02_操作マニュアル.docx')
    doc.add_paragraph('• デザイン文書: platform/docs/design/')
    doc.add_paragraph('• テスト報告: platform/docs/test/')
    
    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph(f'© 2026 Thirdwave AI Platform - README')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.save('/home/nvidia/AI_Coding_Agent/Kadavuley/AI-Coding-Agent/platform/docs/manual/60_03_README.docx')
    print('✓ 60_03_README.docx 生成完了')

def create_60_04_faq():
    """60_04: FAQ (よくある質問 / 頻出QA)"""
    doc = Document()
    
    # Title
    title = doc.add_paragraph()
    title_run = title.add_run('Thirdwave AI Coding Platform\nFAQ / よくある質問')
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f'Frequently Asked Questions / バージョン 0.1.0\n生成日: {datetime.now().strftime("%Y年%m月%d日")}')
    
    # FAQ content
    faqs = [
        {
            'q': 'Q1: Thirdwave とは何ですか？',
            'a': 'Thirdwave は自己ホスト型の AI コーディング支援プラットフォームです。'
                 'ローカル GPU 推論とクラウド AI の両方に対応し、'
                 'Human-in-the-Loop ガバナンスにより安全な自動化を提供します。'
        },
        {
            'q': 'Q2: どのような AI モデルがサポートされていますか？',
            'a': 'ローカル: Qwen3-Coder, MiniMax M2.1 など\n'
                 'クラウド: OpenAI (GPT-4), Anthropic (Claude), Google AI, Mistral など\n'
                 '複数のモデルを同時に利用可能。モデル間での自動フォールバック対応。'
        },
        {
            'q': 'Q3: インストールに GPU は必須ですか？',
            'a': 'いいえ。GPU があると高速ですが、CPU のみでも動作します。\n'
                 'ローカル推論を使わず、クラウド AI のみの利用も可能です。'
        },
        {
            'q': 'Q4: どのポートを使用しますか？',
            'a': 'Platform API: 3100\n'
                 'OpenCode: 4096\n'
                 'nginx: 80/443 (本番)\n'
                 'vLLM: 8000 (GPU サーバー)\n'
                 'これらはすべて設定可能です。'
        },
        {
            'q': 'Q5: マルチユーザーをサポートしていますか？',
            'a': 'はい。THIRDWAVE_PORT_OFFSET を使用して複数インスタンスを実行できます。\n'
                 '例: THIRDWAVE_PORT_OFFSET=10 → ポート 3110, 4106\n'
                 'Docker/systemd で複数サービスを並行運用可能。'
        },
        {
            'q': 'Q6: Human-in-the-Loop (HITL) とは何ですか？',
            'a': 'リスク評価に基づいて自動実行 / ユーザー確認 / 完全ブロック を判断します。\n'
                 '例: 破壊的コマンド (rm -rf) や機密ファイル (.env) へのアクセスは自動ブロック。\n'
                 'セキュリティと利便性のバランスを取ります。'
        },
        {
            'q': 'Q7: セッションはどこに保存されますか？',
            'a': 'SQLite WAL モード: ~/.platform/platform.db\n'
                 'OpenCode エンジン状態: ~/.local/share/opencode/\n'
                 '監査ログ: ~/.platform/audit.db\n'
                 'すべて永続的に保存されるため、再起動後も履歴が保持されます。'
        },
        {
            'q': 'Q8: どのような監査ログが記録されますか？',
            'a': 'セッション作成 / 削除\n'
                 'メッセージ送信\n'
                 'ファイル操作 (読込, 書込, 削除)\n'
                 'コマンド実行\n'
                 'モデル / スキル選択\n'
                 'HITL 承認 / 拒否 決定\n'
                 'すべてタイムスタンプ付きで記録。'
        },
        {
            'q': 'Q9: スキルとは何ですか？',
            'a': '48+ のドメイン別専門ガイドです。\n'
                 '例: API 設計, テスト駆動開発, Azure 最適化, React ベストプラクティス\n'
                 'ユーザーの質問に応じて自動的に関連スキルが RAG コンテキストに注入されます。'
        },
        {
            'q': 'Q10: VS Code 拡張機能はどのようにインストールしますか？',
            'a': '1. VS Code の拡張機能パネルで "Thirdwave AI" を検索\n'
                 '2. インストール & 有効化\n'
                 '3. Platform API が http://localhost:3100 で実行中であることを確認\n'
                 '4. Chat パレットで @thirdwave を使用開始'
        },
        {
            'q': 'Q11: API キーの管理はどうされていますか？',
            'a': 'クラウドプロバイダー API キーは環境変数で管理されます。\n'
                 '例: OPENAI_API_KEY, ANTHROPIC_API_KEY など\n'
                 '機密ファイル保護ポリシーにより .env ファイルアクセスは自動ブロック。'
        },
        {
            'q': 'Q12: Docker でのデプロイはどうしますか？',
            'a': 'docker compose -f platform/docker/docker-compose.yml up\n'
                 '.env ファイルで VLLM_GATEWAY_URL などを設定\n'
                 'opencode-state ボリュームで状態を永続化\n'
                 'Health check により自動再起動対応。'
        },
        {
            'q': 'Q13: nginx でのレート制限は？',
            'a': '一般 API: 30 req/s (バースト 50)\n'
                 'チャット: 5 req/s (バースト 10)\n'
                 'IP あたり 20 同時接続まで\n'
                 'これらは nginx 設定で調整可能。'
        },
        {
            'q': 'Q14: バックアップはどのようにしますか？',
            'a': '重要なデータ:\n'
                 '• ~/.platform/ (SQLite DB, 監査ログ)\n'
                 '• ~/.local/share/opencode/ (セッション状態)\n'
                 '• 設定ファイル (.env)\n'
                 'スクリプトで自動バックアップを推奨。'
        },
        {
            'q': 'Q15: 性能を向上させるには？',
            'a': '1. vLLM キャッシュを有効化 (max_model_len 適切に設定)\n'
                 '2. Redis キャッシング (オプション)\n'
                 '3. マルチワーカーで並列実行\n'
                 '4. nginx バッファリング設定最適化\n'
                 '5. systemd のワーカー数を増やす'
        },
    ]
    
    for idx, faq in enumerate(faqs, 1):
        add_heading_style(doc, faq['q'], 2)
        doc.add_paragraph(faq['a'])
        doc.add_paragraph()
    
    # Additional Resources
    add_heading_style(doc, 'その他の情報 (Additional Resources)', 1)
    doc.add_paragraph('• GitHub リポジトリ: <project-url>')
    doc.add_paragraph('• ドキュメント: platform/docs/')
    doc.add_paragraph('• API リファレンス: http://localhost:3100/health （ドキュメント付き）')
    doc.add_paragraph('• コミュニティフォーラム: <community-link>')
    
    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph(f'© 2026 Thirdwave AI Platform - FAQ')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.save('/home/nvidia/AI_Coding_Agent/Kadavuley/AI-Coding-Agent/platform/docs/manual/60_04_FAQ.docx')
    print('✓ 60_04_FAQ.docx 生成完了')

if __name__ == '__main__':
    print('📄 Thirdwave AI Platform DOCX ドキュメント生成開始...')
    print()
    
    create_60_01_admin_manual()
    create_60_02_user_manual()
    create_60_03_readme()
    create_60_04_faq()
    
    print()
    print('✅ すべてのドキュメント生成完了！')
    print('📁 出力先: platform/docs/manual/')
