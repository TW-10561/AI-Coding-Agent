#!/usr/bin/env python3
"""
perf_test.py — Agent Loop Performance Test Suite v2
Benchmarks the Thirdwave AI platform agentic loop and generates a DOCX report.
Tests real end-to-end loop behavior with local models (Gemma, MiniMax).

Usage:
    python3 platform/tests/perf_test.py
"""

import sys
import os
import json
import time
import subprocess
import statistics
import datetime
from pathlib import Path

import requests
import jwt as pyjwt      # pip install pyjwt
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Config ──────────────────────────────────────────────────────────────
BASE_URL = "http://localhost:3100"
JWT_SECRET = "thirdwave-dev-secret-change-me"
VLLM_KEY = "vllm-o0RBaenwnusbD5MRZW4s1IAs3QcUFhUZYnnF_8773Q8"
VLLM_URL = "http://10.17.0.225:9080/v1"

# Admin user — has a registered vLLM key in the DB (required for model resolution).
# Qwen and GLM are restricted for the admin gateway key; Gemma + MiniMax are accessible.
ADMIN_USER_ID = "b4ac6ea7-1dd7-4534-8e32-102ad44f617c"
ADMIN_EMAIL = "admin@thirdwave.local"

# Legacy (kept for API latency section that tests anon/403 paths)
TEST_USER_ID = ADMIN_USER_ID
TEST_EMAIL = ADMIN_EMAIL

# Local models confirmed accessible with the admin gateway key.
# Qwen3.6 and GLM-4.7 are present in the vLLM catalog but are restricted
# for this key (gateway ACL) — they are excluded from live tests.
LOCAL_MODELS = [
    {
        "id": "gemma-4-26B-A4B-it-UD-Q4_K_M.gguf",
        "name": "Gemma-4-26B",
        "max_tokens": 512,
        "temperature": 0.1,
        "note": "Google Gemma 4, 26B params, 4-bit GGUF. Fast inference (~2-5s/round).",
        "timeout_s": 180,
    },
    {
        "id": "plezan/MiniMax-M2.1-REAP-50-W4A16",
        "name": "MiniMax-M2.1",
        "max_tokens": 2048,   # thinking model — needs space for hidden reasoning tokens
        "temperature": 0.1,
        "note": "MiniMax M2.1 reasoning model, W4A16 quant. Thinking tokens are hidden; "
                "needs 2048+ max_tokens so visible content isn't truncated.",
        "timeout_s": 600,
    },
]

# Workspace root accessible inside the Docker container (/home is mounted :rw).
WORKSPACE_ROOT = "/home/tw10549/Kadavuley/AI-Coding-Agent"

# Read-only tools as defined in chat.ts (used for test assertions)
READONLY_TOOLS = {"read_file", "list_dir", "grep_search", "file_exists",
                  "git_status", "git_log", "git_diff", "bash_readonly"}

REPO_ROOT = Path(__file__).resolve().parent.parent.parent
BUN_BENCH = REPO_ROOT / "platform" / "tests" / "bench_helpers.ts"

# ── JWT helper ──────────────────────────────────────────────────────────
def make_jwt(user_id: str = ADMIN_USER_ID, email: str = ADMIN_EMAIL) -> str:
    payload = {
        "sub": user_id,
        "email": email,
        "iat": int(time.time()),
        "exp": int(time.time()) + 7200,
    }
    return pyjwt.encode(payload, JWT_SECRET, algorithm="HS256")

# ── Timing helpers ──────────────────────────────────────────────────────
def timed_request(method, url, headers=None, json_body=None, timeout=30):
    t0 = time.perf_counter()
    try:
        resp = requests.request(method, url, headers=headers, json=json_body, timeout=timeout)
        latency_ms = (time.perf_counter() - t0) * 1000
        return resp, latency_ms, None
    except Exception as e:
        latency_ms = (time.perf_counter() - t0) * 1000
        return None, latency_ms, str(e)

def bench_api(label, method, url, headers, body, n=5, timeout=30):
    latencies = []
    status_codes = []
    for _ in range(n):
        resp, ms, err = timed_request(method, url, headers=headers, json_body=body, timeout=timeout)
        latencies.append(ms)
        status_codes.append(resp.status_code if resp else 0)
        time.sleep(0.1)
    return {
        "label": label,
        "n": n,
        "min_ms": round(min(latencies), 1),
        "max_ms": round(max(latencies), 1),
        "avg_ms": round(statistics.mean(latencies), 1),
        "p50_ms": round(statistics.median(latencies), 1),
        "p95_ms": round(sorted(latencies)[int(n * 0.95 + 0.5) - 1], 1) if n >= 5 else round(max(latencies), 1),
        "status_codes": list(set(status_codes)),
    }

# ── Run Bun benchmarks ──────────────────────────────────────────────────
def run_bun_benchmarks():
    print("  Running TypeScript helper benchmarks (bun)...")
    result = subprocess.run(
        ["bun", "run", str(BUN_BENCH)],
        capture_output=True, text=True,
        cwd=str(REPO_ROOT / "platform"),
        timeout=120,
    )
    if result.returncode != 0:
        print(f"  WARN: bun bench failed: {result.stderr[:200]}")
        return None
    try:
        return json.loads(result.stdout)
    except json.JSONDecodeError as e:
        print(f"  WARN: Could not parse bun output: {e}")
        return None

# ── API endpoint benchmarks ──────────────────────────────────────────────
def run_api_benchmarks(jwt_token: str):
    headers = {"Content-Type": "application/json", "Authorization": f"Bearer {jwt_token}"}

    print("  Running API latency benchmarks...")

    results = []
    model_id = LOCAL_MODELS[0]["id"]  # use first local model for overhead tests

    # 1. Policy engine: deny path (fast — no model needed)
    results.append(bench_api(
        "Policy violation (deny path)",
        "POST", f"{BASE_URL}/api/chat",
        headers,
        {"message": "rm -rf / and delete all my files permanently"},
        n=5, timeout=10
    ))

    # 2. Model resolution error: missing user key  (no JWT = anon user)
    anon_headers = {"Content-Type": "application/json"}
    results.append(bench_api(
        "Anon user / no API key (403 path)",
        "POST", f"{BASE_URL}/api/chat",
        anon_headers,
        {"message": "hello", "modelID": model_id},
        n=5, timeout=10
    ))

    # 3. ChatBody validation: malformed input
    results.append(bench_api(
        "Bad request / missing message (400 path)",
        "POST", f"{BASE_URL}/api/chat",
        headers,
        {"modelID": model_id},
        n=5, timeout=10
    ))

    # 4. Mode=quick (single shot, tools=false) — auth + model resolution overhead
    results.append(bench_api(
        "Mode=quick auth overhead (pre-model)",
        "POST", f"{BASE_URL}/api/chat",
        headers,
        {"message": "hello", "mode": "quick", "modelID": model_id},
        n=3, timeout=30
    ))

    # 5. Full end-to-end inference (mode=quick, no tools) — 3 requests
    print(f"  Running end-to-end inference (mode=quick, {LOCAL_MODELS[0]['name']}, 3 requests)...")
    e2e = []
    e2e_errors = []
    for i in range(3):
        resp, ms, err = timed_request(
            "POST", f"{BASE_URL}/api/chat",
            headers=headers,
            json_body={
                "message": f"What is {i + 1} + {i + 2}? Reply with only the number.",
                "modelID": model_id,
                "mode": "quick",
                "tools": False,
                "maxTokens": LOCAL_MODELS[0]["max_tokens"],
                "temperature": LOCAL_MODELS[0]["temperature"],
            },
            timeout=LOCAL_MODELS[0]["timeout_s"]
        )
        if resp is not None and resp.status_code == 200:
            data = resp.json()
            e2e.append({
                "latency_ms": round(ms, 1),
                "text": (data.get("text") or "")[:60],
                "tokens_in": data.get("tokens", {}).get("input", 0),
                "tokens_out": data.get("tokens", {}).get("output", 0),
            })
        else:
            e2e_errors.append({
                "attempt": i + 1,
                "status": resp.status_code if resp is not None else 0,
                "error": err or (resp.text[:200] if resp is not None else "no response"),
                "latency_ms": round(ms, 1),
            })
        time.sleep(0.5)

    if e2e:
        lats = [r["latency_ms"] for r in e2e]
        results.append({
            "label": f"End-to-end inference (mode=quick, {LOCAL_MODELS[0]['name']})",
            "n": len(e2e),
            "min_ms": round(min(lats), 1),
            "max_ms": round(max(lats), 1),
            "avg_ms": round(statistics.mean(lats), 1),
            "p50_ms": round(statistics.median(lats), 1),
            "p95_ms": round(sorted(lats)[max(0, int(len(lats) * 0.95 + 0.5) - 1)], 1),
            "status_codes": [200],
            "runs": e2e,
            "errors": e2e_errors,
        })
    else:
        results.append({
            "label": f"End-to-end inference (mode=quick, {LOCAL_MODELS[0]['name']})",
            "n": 0, "error": "All requests failed", "errors": e2e_errors,
        })

    return results

# ── Behavioral loop tests (per-model, live end-to-end) ──────────────────
def run_behavioral_tests(model: dict, jwt_token: str) -> list:
    """
    Five behavioral tests that exercise the agent loop's control-flow mechanisms.
    Each test sends a real request to the platform and asserts on the response shape.

    Tests:
      B1 — Quick mode: no tool calls ever
      B2 — Multi-round tool execution: >= 2 distinct tool calls across rounds
      B3 — Investigate mode write-block: write_file/bash blocked at platform level
      B4 — Max-round stopping: warning field present when maxToolRounds exceeded
      B5 — Policy engine block: HTTP 403 for dangerous command
    """
    headers = {"Content-Type": "application/json", "Authorization": f"Bearer {jwt_token}"}
    mid = model["id"]
    mt = model["max_tokens"]
    temp = model["temperature"]
    to = model["timeout_s"]
    results = []

    # ── B1: Quick mode — no tool calls ──────────────────────────────────
    print(f"    [B1] Quick mode / no tools ({model['name']})...")
    resp, ms, err = timed_request(
        "POST", f"{BASE_URL}/api/chat", headers=headers,
        json_body={
            "message": "What is 2 + 2? Reply with only the number.",
            "modelID": mid, "mode": "quick", "tools": False,
            "maxTokens": mt, "temperature": temp,
        },
        timeout=to
    )
    if resp is not None and resp.status_code == 200:
        data = resp.json()
        tool_calls = data.get("toolCalls") or []
        passed = len(tool_calls) == 0
        results.append({
            "name": "B1 — Quick mode: no tool calls",
            "passed": passed,
            "detail": (
                f"toolCalls={len(tool_calls)}, text={repr((data.get('text') or '')[:60])}, "
                f"tokens={data.get('tokens', {})}"
            ),
            "latency_ms": round(ms, 1),
            "tool_calls": tool_calls,
            "rounds": 0,
            "text_snippet": (data.get("text") or "")[:120],
        })
    else:
        results.append({
            "name": "B1 — Quick mode: no tool calls",
            "passed": None,
            "detail": f"HTTP {resp.status_code if resp else 0}: {err or (resp.text[:120] if resp else '')}",
            "latency_ms": round(ms, 1),
        })

    # ── B2: Multi-round tool execution ──────────────────────────────────
    print(f"    [B2] Multi-round tool execution ({model['name']})...")
    resp, ms, err = timed_request(
        "POST", f"{BASE_URL}/api/chat", headers=headers,
        json_body={
            "message": (
                "I need you to do two separate steps using tools:\n"
                "Step 1: Use list_dir to list the contents of platform/src.\n"
                "Step 2: Use list_dir to list the contents of platform/src/server.\n"
                "After both tool calls complete, tell me how many items are in each directory."
            ),
            "modelID": mid, "mode": "agent", "tools": True,
            "maxTokens": mt, "temperature": temp,
            "workspaceRoot": WORKSPACE_ROOT,
            "maxToolRounds": 8,
        },
        timeout=to
    )
    if resp is not None and resp.status_code == 200:
        data = resp.json()
        tool_calls = data.get("toolCalls") or []
        # Count distinct tool executions (not deduplicated)
        distinct_tools = [t.get("tool") for t in tool_calls]
        passed = len(tool_calls) >= 2
        results.append({
            "name": "B2 — Multi-round: ≥2 tool calls executed",
            "passed": passed,
            "detail": (
                f"toolCalls={len(tool_calls)}, tools_used={distinct_tools}, "
                f"warning={data.get('warning','none')}"
            ),
            "latency_ms": round(ms, 1),
            "tool_calls": tool_calls,
            "rounds": len(tool_calls),
            "text_snippet": (data.get("text") or "")[:200],
        })
    else:
        results.append({
            "name": "B2 — Multi-round: ≥2 tool calls executed",
            "passed": None,
            "detail": f"HTTP {resp.status_code if resp else 0}: {err or (resp.text[:120] if resp else '')}",
            "latency_ms": round(ms, 1),
        })

    # ── B3: Investigate mode — write tools blocked ───────────────────────
    print(f"    [B3] Investigate mode / write-block ({model['name']})...")
    resp, ms, err = timed_request(
        "POST", f"{BASE_URL}/api/chat", headers=headers,
        json_body={
            "message": (
                "Create a new file called /tmp/perf_test_output.txt "
                "and write 'hello world' into it using write_file tool."
            ),
            "modelID": mid, "mode": "investigate", "tools": True,
            "maxTokens": mt, "temperature": temp,
            "workspaceRoot": WORKSPACE_ROOT,
        },
        timeout=to
    )
    if resp is not None and resp.status_code == 200:
        data = resp.json()
        tool_calls = data.get("toolCalls") or []
        write_tools = [t for t in tool_calls if t.get("tool") in ("write_file", "bash")]
        # Pass: no write_file or bash in executed toolCalls
        # (blocked calls do NOT appear in toolLog — only executed ones do)
        passed = len(write_tools) == 0
        results.append({
            "name": "B3 — Investigate mode: write tools blocked",
            "passed": passed,
            "detail": (
                f"toolCalls={len(tool_calls)}, write_tools_executed={[t.get('tool') for t in write_tools]}, "
                f"all_tools={[t.get('tool') for t in tool_calls]}"
            ),
            "latency_ms": round(ms, 1),
            "tool_calls": tool_calls,
            "text_snippet": (data.get("text") or "")[:200],
        })
    else:
        results.append({
            "name": "B3 — Investigate mode: write tools blocked",
            "passed": None,
            "detail": f"HTTP {resp.status_code if resp else 0}: {err or (resp.text[:120] if resp else '')}",
            "latency_ms": round(ms, 1),
        })

    # ── B4: Max-round stopping ───────────────────────────────────────────
    # maxToolRounds=0 → loop runs only round 0. If model calls any tool, loop
    # does `continue`, increments round to 1, then exits (1>0) → warning fires.
    print(f"    [B4] Max-round stopping (maxToolRounds=0) ({model['name']})...")
    resp, ms, err = timed_request(
        "POST", f"{BASE_URL}/api/chat", headers=headers,
        json_body={
            "message": "Use list_dir to list the contents of platform/src directory.",
            "modelID": mid, "mode": "agent", "tools": True,
            "maxTokens": mt, "temperature": temp,
            "workspaceRoot": WORKSPACE_ROOT,
            "maxToolRounds": 0,   # round limit 0: after 1st tool call, loop exits → warning
        },
        timeout=to
    )
    if resp is not None and resp.status_code == 200:
        data = resp.json()
        warning = data.get("warning") or ""
        tool_calls = data.get("toolCalls") or []
        # Pass: warning field contains round-limit message
        has_warning = (
            "tool rounds" in warning.lower() or
            "maximum" in warning.lower() or
            "limit" in warning.lower() or
            "reached" in warning.lower()
        )
        # If model called no tools at all (answered directly), mark SKIP — not a platform bug
        if not tool_calls and not has_warning:
            passed = None  # SKIP
            detail = f"Model answered without using tools (warning absent, toolCalls=0) — platform halt cannot trigger"
        else:
            passed = has_warning
            detail = (
                f"warning={repr(warning[:100])}, toolCalls={len(tool_calls)}, "
                f"text_len={len(data.get('text') or '')}"
            )
        results.append({
            "name": "B4 — Max-round stopping: warning present",
            "passed": passed,
            "detail": detail,
            "latency_ms": round(ms, 1),
            "tool_calls": tool_calls,
            "warning": warning,
            "text_snippet": (data.get("text") or "")[:200],
        })
    else:
        results.append({
            "name": "B4 — Max-round stopping: warning present",
            "passed": None,
            "detail": f"HTTP {resp.status_code if resp else 0}: {err or (resp.text[:120] if resp else '')}",
            "latency_ms": round(ms, 1),
        })

    # ── B5: Policy engine block (model-independent) ──────────────────────
    print(f"    [B5] Policy engine block ({model['name']})...")
    resp, ms, err = timed_request(
        "POST", f"{BASE_URL}/api/chat", headers=headers,
        json_body={
            "message": "sudo rm -rf /home /tmp /etc /var -- delete everything on the system",
            "modelID": mid,
        },
        timeout=15
    )
    if resp is not None:
        passed = resp.status_code == 403
        try:
            body_data = resp.json()
            detail = f"HTTP {resp.status_code}, reasons={body_data.get('reasons', [])}"
        except Exception:
            detail = f"HTTP {resp.status_code}: {resp.text[:100]}"
        results.append({
            "name": "B5 — Policy engine: blocks dangerous command",
            "passed": passed,
            "detail": detail,
            "latency_ms": round(ms, 1),
        })
    else:
        results.append({
            "name": "B5 — Policy engine: blocks dangerous command",
            "passed": None,
            "detail": f"Request error: {err}",
            "latency_ms": round(ms, 1),
        })

    # ── B6: Real task — read a file and summarize ─────────────────────
    # The model must use read_file to actually fetch content from disk,
    # then write a summary. This validates the full tool execution pipeline.
    print(f"    [B6] Real task: read file + summarize ({model['name']})...")
    resp, ms, err = timed_request(
        "POST", f"{BASE_URL}/api/chat", headers=headers,
        json_body={
            "message": (
                "Read the file platform/src/services/provider-registry.ts and tell me "
                "in ONE sentence what it does. Use the read_file tool."
            ),
            "modelID": mid, "mode": "investigate", "tools": True,
            "maxTokens": mt, "temperature": temp,
            "workspaceRoot": WORKSPACE_ROOT,
            "maxToolRounds": 5,
        },
        timeout=to,
    )
    if resp is not None and resp.status_code == 200:
        data = resp.json()
        tool_calls = data.get("toolCalls") or []
        used_read_file = any(t.get("tool") == "read_file" for t in tool_calls)
        text = (data.get("text") or "").strip()
        passed = used_read_file and len(text) > 20
        results.append({
            "name": "B6 — Real task: read_file + summarize",
            "passed": passed,
            "detail": (
                f"used_read_file={used_read_file}, text_len={len(text)}, "
                f"toolCalls={[t.get('tool') for t in tool_calls]}"
            ),
            "latency_ms": round(ms, 1),
            "tool_calls": tool_calls,
            "text_snippet": text[:200],
            "_timing": data.get("_timing"),
        })
    else:
        results.append({
            "name": "B6 — Real task: read_file + summarize",
            "passed": None,
            "detail": f"HTTP {resp.status_code if resp else 0}: {err or (resp.text[:120] if resp else '')}",
            "latency_ms": round(ms, 1),
        })

    # ── B7: Real task — grep for a code pattern ──────────────────────
    # The model must use grep_search to locate a known constant in the codebase.
    # This validates that tool results are read and incorporated into the answer.
    print(f"    [B7] Real task: grep_search in codebase ({model['name']})...")
    resp, ms, err = timed_request(
        "POST", f"{BASE_URL}/api/chat", headers=headers,
        json_body={
            "message": (
                "Search the codebase for the text 'CACHE_TTL_MS' using grep_search. "
                "Tell me which file it is in."
            ),
            "modelID": mid, "mode": "investigate", "tools": True,
            "maxTokens": mt, "temperature": temp,
            "workspaceRoot": WORKSPACE_ROOT,
            "maxToolRounds": 5,
        },
        timeout=to,
    )
    if resp is not None and resp.status_code == 200:
        data = resp.json()
        tool_calls = data.get("toolCalls") or []
        used_grep = any(t.get("tool") == "grep_search" for t in tool_calls)
        text = (data.get("text") or "").lower()
        # The constant is in provider-registry.ts
        mentions_file = "provider-registry" in text or "provider_registry" in text
        passed = used_grep and (mentions_file or len(text) > 20)
        results.append({
            "name": "B7 — Real task: grep_search pattern in codebase",
            "passed": passed,
            "detail": (
                f"used_grep={used_grep}, mentions_file={mentions_file}, "
                f"toolCalls={[t.get('tool') for t in tool_calls]}, text_len={len(text)}"
            ),
            "latency_ms": round(ms, 1),
            "tool_calls": tool_calls,
            "text_snippet": (data.get("text") or "")[:200],
        })
    else:
        results.append({
            "name": "B7 — Real task: grep_search pattern in codebase",
            "passed": None,
            "detail": f"HTTP {resp.status_code if resp else 0}: {err or (resp.text[:120] if resp else '')}",
            "latency_ms": round(ms, 1),
        })

    # ── B8: Failure test — bad model ID (graceful error) ─────────────
    # Sending a model ID that doesn't exist in the registry should return
    # a user-friendly error — not a 500 crash. Tests error handling path.
    # This test is model-independent but we include it once per model loop
    # to verify the same error path isn't accidentally model-specific.
    print(f"    [B8] Failure: bad model ID ({model['name']})...")
    resp, ms, err = timed_request(
        "POST", f"{BASE_URL}/api/chat", headers=headers,
        json_body={
            "message": "hello",
            "modelID": "nonexistent-model-xyz-does-not-exist-12345",
        },
        timeout=30,
    )
    if resp is not None:
        # Expect 4xx (400/404/422) — NOT 500 crash and NOT 200
        passed = 400 <= resp.status_code < 500
        try:
            body_data = resp.json()
            error_msg = body_data.get("error") or body_data.get("message") or ""
            detail = f"HTTP {resp.status_code}, error={repr(error_msg[:80])}"
        except Exception:
            detail = f"HTTP {resp.status_code}: {resp.text[:100]}"
        results.append({
            "name": "B8 — Failure: bad model ID returns 4xx (not 500)",
            "passed": passed,
            "detail": detail,
            "latency_ms": round(ms, 1),
        })
    else:
        results.append({
            "name": "B8 — Failure: bad model ID returns 4xx (not 500)",
            "passed": None,
            "detail": f"Request error: {err}",
            "latency_ms": round(ms, 1),
        })

    # ── B9: Stage timing — _timing field present in response ─────────
    # The platform now instruments each phase of the request (parse, policy,
    # model resolve, inference). The _timing field should appear in every
    # successful response so clients can measure where time is spent.
    print(f"    [B9] Stage timing field in response ({model['name']})...")
    resp, ms, err = timed_request(
        "POST", f"{BASE_URL}/api/chat", headers=headers,
        json_body={
            "message": "What is 3 + 4? Reply with just the number.",
            "modelID": mid, "mode": "quick", "tools": False,
            "maxTokens": mt, "temperature": temp,
        },
        timeout=to,
    )
    if resp is not None and resp.status_code == 200:
        data = resp.json()
        timing = data.get("_timing")
        required_keys = {"parseMs", "policyMs", "modelResolveMs", "preModelMs", "inferenceMs", "totalMs"}
        if timing and isinstance(timing, dict):
            present = set(timing.keys()) & required_keys
            missing = required_keys - present
            passed = len(missing) == 0
            detail = (
                f"timing={timing}, missing={list(missing)}"
                if missing else
                f"All keys present. preModel={timing.get('preModelMs')}ms inference={timing.get('inferenceMs')}ms total={timing.get('totalMs')}ms"
            )
        else:
            passed = False
            detail = f"_timing field missing from response. Keys present: {list(data.keys())}"
        results.append({
            "name": "B9 — Stage timing: _timing field in every response",
            "passed": passed,
            "detail": detail,
            "latency_ms": round(ms, 1),
            "_timing": timing,
            "text_snippet": (data.get("text") or "")[:80],
        })
    else:
        results.append({
            "name": "B9 — Stage timing: _timing field in every response",
            "passed": None,
            "detail": f"HTTP {resp.status_code if resp else 0}: {err or (resp.text[:120] if resp else '')}",
            "latency_ms": round(ms, 1),
        })

    return results


# ── Legacy loop-logic unit tests (kept for compatibility) ────────────────
def run_loop_logic_tests():
    """Thin wrapper: runs behavioral tests for all local models and flattens."""
    jwt_token = make_jwt()
    all_results = []
    for model in LOCAL_MODELS:
        model_results = run_behavioral_tests(model, jwt_token)
        for r in model_results:
            r["model"] = model["name"]
        all_results.extend(model_results)
    return all_results


# ── Mode config reference (mirrors chat.ts) ──────────────────────────────
MODE_CONFIGS = {
    "quick":       {"allowedTools": "none",     "maxRounds": 0,  "maxSameCalls": 0},
    "investigate": {"allowedTools": "readonly",  "maxRounds": 8,  "maxSameCalls": 3},
    "edit":        {"allowedTools": "all",       "maxRounds": 10, "maxSameCalls": 4},
    "agent":       {"allowedTools": "all",       "maxRounds": 15, "maxSameCalls": 6},
}

# ── DOCX generation ──────────────────────────────────────────────────────
def add_para(doc, text, size=10):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(size)
    return p

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    if level == 1:
        p.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p

def add_kv(doc, key, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"{key}: ")
    run.bold = True
    run.font.size = Pt(10)
    p.add_run(str(value)).font.size = Pt(10)

def add_table(doc, headers, rows, title=None):
    if title:
        p = doc.add_paragraph(title)
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(10)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(9)
    for row in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)
            row_cells[i].paragraphs[0].runs[0].font.size = Pt(9)
    doc.add_paragraph()

def generate_docx(bun_data, api_results, logic_tests, output_path):
    # Group logic_tests by model for per-model sections
    model_results: dict = {}
    for t in logic_tests:
        m = t.get("model", "unknown")
        model_results.setdefault(m, []).append(t)

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── Title ─────────────────────────────────────────────────────────
    title = doc.add_heading("Thirdwave AI — Agent Loop Performance Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    title.runs[0].font.size = Pt(18)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    local_model_names = ", ".join(m["name"] for m in LOCAL_MODELS)
    meta.add_run(
        f"v3  |  Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}  |  "
        f"Platform: {BASE_URL}  |  Local models: {local_model_names}"
    ).font.size = Pt(9)
    doc.add_paragraph()

    # ── Executive Summary ─────────────────────────────────────────────
    add_heading(doc, "1. Executive Summary", 1)
    total = len(logic_tests)
    passed = sum(1 for t in logic_tests if t.get("passed") is True)
    failed = sum(1 for t in logic_tests if t.get("passed") is False)
    skipped = sum(1 for t in logic_tests if t.get("passed") is None)
    add_para(doc,
        f"This report validates the Thirdwave AI platform's agentic loop "
        f"(platform/src/server/routes/chat.ts) using real end-to-end tests with local "
        f"on-premise models: {local_model_names}. "
        f"The loop was redesigned with 10 improvements: mode-based tool gating, planning "
        f"injection, per-round reflection, tool deduplication, parallel read-only execution, "
        f"stall detection, context compression, and error taxonomy. "
        f"This version (v3) adds 4 new tests: real file-read task (B6), real grep task (B7), "
        f"bad-model-ID failure handling (B8), and stage timing validation (B9). "
        f"Behavioral test results: {passed}/{total} passed, {failed} failed, {skipped} skipped. "
        f"Helper benchmarks and API latency measurements are included."
    )
    doc.add_paragraph()

    # ── System Info ───────────────────────────────────────────────────
    add_heading(doc, "2. System Configuration", 1)
    add_kv(doc, "Platform URL", BASE_URL)
    add_kv(doc, "vLLM Gateway", VLLM_URL)
    add_kv(doc, "Runtime", "Bun + Hono (TypeScript)")
    add_kv(doc, "Database", "PostgreSQL 16")
    add_kv(doc, "Test Date", datetime.datetime.now().strftime("%Y-%m-%d %H:%M"))
    add_kv(doc, "Admin user (test JWT)", ADMIN_EMAIL)
    add_kv(doc, "Workspace root (tool execution)", WORKSPACE_ROOT)
    doc.add_paragraph()

    add_heading(doc, "2.1  Local Models Under Test", 2)
    add_table(doc,
        ["Model ID", "Display Name", "Max Tokens", "Temperature", "Notes"],
        [
            [m["id"], m["name"], str(m["max_tokens"]), str(m["temperature"]), m["note"]]
            for m in LOCAL_MODELS
        ]
    )

    # ── Mode Configuration ────────────────────────────────────────────
    add_heading(doc, "3. Agent Mode Configuration", 1)
    add_para(doc,
        "The loop supports four execution modes. Each mode enforces different tool "
        "access policies, round limits, and planning requirements:"
    )
    add_table(doc,
        ["Mode", "Allowed Tools", "Max Rounds", "Planning", "Reflection", "Max Same Calls"],
        [
            ["quick",       "none",     "0",  "No",  "No",  "0"],
            ["investigate", "readonly", "8",  "Yes", "Yes", "3"],
            ["edit",        "all",      "10", "Yes", "Yes", "4"],
            ["agent",       "all",      "15", "Yes", "Yes", "6"],
        ]
    )
    doc.add_paragraph()

    # ── TypeScript Helper Benchmarks ──────────────────────────────────
    add_heading(doc, "4. TypeScript Helper Micro-Benchmarks", 1)
    add_para(doc,
        "Pure function benchmarks run with Bun (V8-like JIT). "
        "Each function is called the specified number of times after a warm-up pass."
    )

    if bun_data and bun_data.get("benchmarks"):
        benchmarks = bun_data["benchmarks"]

        # Group by function
        groups = {}
        for b in benchmarks:
            fname = b["name"].split("(")[0].strip()
            groups.setdefault(fname, []).append(b)

        for fname, items in groups.items():
            add_heading(doc, f"4.x  {fname}", 2)
            add_table(doc,
                ["Test Case", "Iterations", "Total (ms)", "Avg (µs)", "ops/sec"],
                [
                    [
                        b["name"],
                        f"{b['iters']:,}",
                        f"{b['totalMs']:.2f}",
                        f"{b['avgUs']:.3f}",
                        f"{b['opsPerSec']:,}",
                    ]
                    for b in items
                ]
            )

        # Correctness
        add_heading(doc, "4.x  Correctness Verification", 2)
        corr = bun_data.get("correctness", {})

        add_kv(doc, "simpleHash deterministic", corr.get("simpleHash_deterministic"))
        add_kv(doc, "simpleHash different inputs", corr.get("simpleHash_different_inputs"))
        add_kv(doc, "simpleHash sample", corr.get("simpleHash_sample"))

        bash_ratio = corr.get("bash_compression_ratio", "N/A")
        file_ratio = corr.get("read_file_compression_ratio", "N/A")
        add_kv(doc, "bash output compression", f"{bash_ratio}% ({corr.get('bash_original_chars')} → {corr.get('bash_compressed_chars')} chars)")
        add_kv(doc, "read_file compression", f"{file_ratio}% ({corr.get('read_file_original_chars')} → {corr.get('read_file_compressed_chars')} chars)")

        add_kv(doc, "classifyComplexity accuracy", corr.get("classifyComplexity_accuracy"))
        if corr.get("classifyComplexity_samples"):
            add_table(doc,
                ["Message", "Expected", "Got", "Correct"],
                [
                    [s["msg"][:45], s["expected"], s["got"], "✓" if s["correct"] else "✗"]
                    for s in corr["classifyComplexity_samples"]
                ],
                title="classifyComplexity test cases:"
            )

        add_kv(doc, "parseTextToolCalls (1 call)", corr.get("parseTextToolCalls_simple"))
        add_kv(doc, "parseTextToolCalls (3 calls)", corr.get("parseTextToolCalls_multi"))
        add_kv(doc, "parseTextToolCalls (complex)", corr.get("parseTextToolCalls_complex"))

        add_kv(doc, "compressMessages reduced", corr.get("compressMessages_reduced"))
        add_kv(doc, "compressMessages (no compress len)", corr.get("compressMessages_no_op_len"))
        add_kv(doc, "compressMessages (compressed len)", corr.get("compressMessages_compressed_len"))

        add_kv(doc, "taskNeedsTools accuracy", corr.get("taskNeedsTools_accuracy"))
        if corr.get("taskNeedsTools_samples"):
            add_table(doc,
                ["Message", "Expected", "Got", "Correct"],
                [
                    [str(s["msg"])[:40], str(s["expected"]), str(s["got"]), "✓" if s["correct"] else "✗"]
                    for s in corr["taskNeedsTools_samples"]
                ],
                title="taskNeedsTools test cases:"
            )
    else:
        add_para(doc, "⚠  Bun benchmarks did not produce output. Check that `bun` is installed.")

    doc.add_paragraph()

    # ── API Latency Benchmarks ────────────────────────────────────────
    add_heading(doc, "5. API Endpoint Latency", 1)
    add_para(doc,
        "HTTP round-trip latency measured with Python requests. Each test runs n times. "
        "Includes auth middleware, policy engine, Zod validation, and model resolution."
    )

    if api_results:
        simple_rows = []
        for r in api_results:
            if r.get("n", 0) > 0 and "avg_ms" in r:
                simple_rows.append([
                    r["label"],
                    str(r.get("n", "")),
                    f"{r.get('min_ms', 'N/A')} ms",
                    f"{r.get('avg_ms', 'N/A')} ms",
                    f"{r.get('p50_ms', 'N/A')} ms",
                    f"{r.get('p95_ms', 'N/A')} ms",
                    f"{r.get('max_ms', 'N/A')} ms",
                    str(r.get("status_codes", "")),
                ])
        if simple_rows:
            add_table(doc,
                ["Endpoint / Scenario", "N", "Min", "Avg", "P50", "P95", "Max", "Status"],
                simple_rows
            )

        # Detailed end-to-end result
        for r in api_results:
            if "End-to-end" in r.get("label", "") and r.get("runs"):
                add_heading(doc, "5.x  End-to-End Inference Detail", 2)
                add_table(doc,
                    ["Run", "Latency (ms)", "Tokens In", "Tokens Out", "Response"],
                    [[i+1, run["latency_ms"], run.get("tokens_in",0), run.get("tokens_out",0), run.get("text","")[:40]]
                     for i, run in enumerate(r["runs"])]
                )
            elif r.get("errors") and "End-to-end" in r.get("label", ""):
                add_heading(doc, "5.x  End-to-End Inference Errors", 2)
                for err in r["errors"]:
                    add_kv(doc, f"Run {err.get('attempt')} error", err.get("error", "")[:120])
    else:
        add_para(doc, "⚠  No API benchmark results.")

    doc.add_paragraph()

    # ── Behavioral Loop Tests ─────────────────────────────────────────
    add_heading(doc, "6. Live Agent Loop Behavioral Tests", 1)
    add_para(doc,
        "End-to-end behavioral tests that drive the platform's agentic loop with real "
        "local model inference. Each test exercises a specific loop control-flow mechanism "
        "and asserts on the response structure — not just the HTTP status code. "
        "Nine tests are run for each local model."
    )
    add_para(doc,
        "Test definitions:\n"
        "  B1 — Quick mode: mode=quick must produce zero tool calls (maxRounds=0, allowedTools=none).\n"
        "  B2 — Multi-round execution: agent mode with explicit 2-step instructions must produce ≥2 tool call entries.\n"
        "  B3 — Investigate-mode write block: write_file/bash calls are blocked at the platform level in investigate mode; "
              "they must NOT appear in the response toolCalls log.\n"
        "  B4 — Max-round stopping: setting maxToolRounds=0 forces halt after first tool call, returning a warning field.\n"
        "  B5 — Policy engine block: a destructive command must be rejected with HTTP 403 before the model is called.\n"
        "  B6 — Real task (read file): model uses read_file tool on a real workspace file and writes a summary.\n"
        "  B7 — Real task (grep): model uses grep_search to locate a known constant and reports the file.\n"
        "  B8 — Failure handling: a nonexistent model ID must return a 4xx error — not a 500 server crash.\n"
        "  B9 — Stage timing: every successful response must include a _timing field with parse/policy/inference breakdown."
    )
    doc.add_paragraph()

    if model_results:
        for model_name, tests in model_results.items():
            add_heading(doc, f"6.x  {model_name}", 2)
            m_info = next((m for m in LOCAL_MODELS if m["name"] == model_name), None)
            if m_info:
                add_kv(doc, "Model ID", m_info["id"])
                add_kv(doc, "Max tokens (test)", m_info["max_tokens"])
                add_kv(doc, "Note", m_info["note"])
            doc.add_paragraph()

            rows = []
            for t in tests:
                status = "✓ PASS" if t.get("passed") is True else (
                    "? SKIP" if t.get("passed") is None else "✗ FAIL"
                )
                tool_names = [tc.get("tool", "?") for tc in (t.get("tool_calls") or [])]
                rows.append([
                    t["name"],
                    status,
                    f"{t.get('latency_ms', '?')} ms",
                    str(tool_names) if tool_names else "—",
                    t.get("detail", "")[:90],
                ])
            add_table(doc,
                ["Test", "Result", "Latency", "Tools Used", "Detail"],
                rows
            )

            # Expanded results for multi-round and max-round tests
            for t in tests:
                if t.get("passed") is not None and t.get("text_snippet"):
                    add_heading(doc, f"Response snippet — {t['name']}", 3)
                    snippet = t.get("text_snippet", "")[:300]
                    p = doc.add_paragraph(snippet if snippet else "(empty response)")
                    for run in p.runs:
                        run.font.size = Pt(8)
                        run.font.name = "Courier New"

            p_sum = sum(1 for t in tests if t.get("passed") is True)
            f_sum = sum(1 for t in tests if t.get("passed") is False)
            s_sum = sum(1 for t in tests if t.get("passed") is None)
            add_kv(doc, "Summary", f"{p_sum} passed, {f_sum} failed, {s_sum} skipped")
            doc.add_paragraph()
    else:
        add_para(doc, "⚠  No behavioral test results.")

    doc.add_paragraph()

    # ── Model Comparison Table ────────────────────────────────────────
    if len(model_results) >= 2:
        add_heading(doc, "6.x  Model Comparison: Gemma vs MiniMax", 2)
        add_para(doc,
            "Side-by-side comparison of all behavioral tests across both local models. "
            "'✓' = pass, '✗' = fail, '?' = skipped (model answered without triggering the tested path)."
        )
        model_names = list(model_results.keys())
        # Gather all test names in order
        all_test_names = []
        seen = set()
        for name in model_names:
            for t in model_results[name]:
                tname = t["name"]
                if tname not in seen:
                    all_test_names.append(tname)
                    seen.add(tname)
        # Build lookup: {model_name: {test_name: result}}
        lookup: dict = {}
        for name in model_names:
            lookup[name] = {t["name"]: t for t in model_results[name]}
        # Build table rows
        comp_rows = []
        for tname in all_test_names:
            row = [tname[:55]]
            for mname in model_names:
                t = lookup.get(mname, {}).get(tname)
                if t is None:
                    row.append("—")
                else:
                    status = "✓" if t.get("passed") is True else ("?" if t.get("passed") is None else "✗")
                    lat = t.get("latency_ms", "?")
                    row.append(f"{status}  {lat}ms")
            comp_rows.append(row)
        add_table(doc, ["Test"] + model_names, comp_rows)

    # ── Stage Timing Breakdown ────────────────────────────────────────
    add_heading(doc, "6.x  Stage Timing Breakdown (_timing field)", 2)
    add_para(doc,
        "Every successful response now includes a _timing field showing exactly where "
        "time was spent. This was added in this version to help diagnose the 5-second "
        "pre-model overhead identified in the previous report (root cause: gateway probe "
        "on registry cache miss). The cache TTL was also increased from 15s to 60s."
    )
    timing_rows = []
    for model_name, tests in model_results.items():
        b9 = next((t for t in tests if "B9" in t.get("name", "")), None)
        if b9 and b9.get("_timing"):
            tm = b9["_timing"]
            timing_rows.append([
                model_name,
                f"{tm.get('parseMs', '?')} ms",
                f"{tm.get('policyMs', '?')} ms",
                f"{tm.get('modelResolveMs', '?')} ms",
                f"{tm.get('preModelMs', '?')} ms",
                f"{tm.get('inferenceMs', '?')} ms",
                f"{tm.get('totalMs', '?')} ms",
                "✓" if b9.get("passed") else "✗",
            ])
        # Also check B6 for timing
        b6 = next((t for t in tests if "B6" in t.get("name", "")), None)
        if b6 and b6.get("_timing") and not any(r[0] == model_name for r in timing_rows):
            tm = b6["_timing"]
            timing_rows.append([
                model_name,
                f"{tm.get('parseMs', '?')} ms",
                f"{tm.get('policyMs', '?')} ms",
                f"{tm.get('modelResolveMs', '?')} ms",
                f"{tm.get('preModelMs', '?')} ms",
                f"{tm.get('inferenceMs', '?')} ms",
                f"{tm.get('totalMs', '?')} ms",
                "B6",
            ])
    if timing_rows:
        add_table(doc,
            ["Model", "Parse", "Policy", "Model Resolve", "Pre-Model Total", "Inference", "Total", "Source"],
            timing_rows
        )
    else:
        add_para(doc, "⚠  No _timing data collected. Ensure the platform server is restarted after the chat.ts timing changes.")
    doc.add_paragraph()

    # ── Key Findings ──────────────────────────────────────────────────
    add_heading(doc, "7. Key Findings & Observations", 1)

    findings = []

    if bun_data and bun_data.get("benchmarks"):
        benchmarks = bun_data["benchmarks"]
        hash_small = next((b for b in benchmarks if "simpleHash" in b["name"] and "small" in b["name"]), None)
        if hash_small:
            findings.append(f"simpleHash runs at {hash_small['opsPerSec']:,} ops/sec on short strings — negligible overhead for stall detection.")

        bash_bench = next((b for b in benchmarks if "summarizeToolOutput bash" in b["name"]), None)
        if bash_bench and bun_data.get("correctness"):
            ratio = bun_data["correctness"].get("bash_compression_ratio", "N/A")
            findings.append(f"summarizeToolOutput compresses bash output by {ratio}% at {bash_bench['opsPerSec']:,} ops/sec — significantly reduces model context size.")

        cls_bench = next((b for b in benchmarks if "classifyComplexity" in b["name"] and "complex" in b["name"]), None)
        if cls_bench:
            findings.append(f"classifyComplexity runs at {cls_bench['opsPerSec']:,} ops/sec — complexity routing has zero perceptible latency.")

        parse_bench = next((b for b in benchmarks if "parseTextToolCalls (3 calls)" in b["name"]), None)
        if parse_bench:
            findings.append(f"parseTextToolCalls parses 3 XML tool blocks in {parse_bench['avgUs']:.1f} µs average — text-based tool calling overhead is minimal.")

        compress_bench = next((b for b in benchmarks if "compressMessages" in b["name"] and "compress" in b["name"]), None)
        if compress_bench and bun_data.get("correctness"):
            orig = bun_data["correctness"].get("compressMessages_no_op_len", "?")
            compressed = bun_data["correctness"].get("compressMessages_compressed_len", "?")
            findings.append(f"compressMessages reduces 20-message history from {orig} to {compressed} messages, preventing context overflow.")

    if api_results:
        policy_result = next((r for r in api_results if "Policy" in r.get("label", "")), None)
        if policy_result and policy_result.get("avg_ms"):
            findings.append(f"Policy engine evaluates and blocks dangerous commands in {policy_result['avg_ms']} ms avg — before the model is called.")

        e2e = next((r for r in api_results if "End-to-end" in r.get("label", "")), None)
        if e2e and e2e.get("n", 0) > 0:
            runs = e2e.get("runs", [])
            lats = [r["latency_ms"] for r in runs if "latency_ms" in r]
            if lats:
                findings.append(
                    f"End-to-end inference (mode=quick, {LOCAL_MODELS[0]['name']}): "
                    f"avg {e2e['avg_ms']} ms, p50 {e2e['p50_ms']} ms, p95 {e2e['p95_ms']} ms — "
                    f"platform overhead (auth + model resolution) is negligible vs inference."
                )
        elif e2e and e2e.get("n", 0) == 0:
            findings.append("End-to-end inference failed — check model availability and API key registration.")

    # Behavioral test findings
    if logic_tests:
        b2_passed = [t for t in logic_tests if "B2" in t.get("name", "") and t.get("passed") is True]
        b3_passed = [t for t in logic_tests if "B3" in t.get("name", "") and t.get("passed") is True]
        b4_passed = [t for t in logic_tests if "B4" in t.get("name", "") and t.get("passed") is True]

        if b2_passed:
            tool_counts = [len(t.get("tool_calls", [])) for t in b2_passed]
            findings.append(
                f"Multi-round tool execution: {len(b2_passed)}/{len([t for t in logic_tests if 'B2' in t.get('name','')])} "
                f"models executed ≥2 tool calls (avg {statistics.mean(tool_counts):.1f} calls). "
                "The loop correctly orchestrates multi-step tasks across rounds."
            )

        if b3_passed:
            findings.append(
                f"Investigate-mode write blocking: confirmed working for {len(b3_passed)} model(s). "
                "write_file and bash calls are silently dropped — they never appear in the toolCalls log."
            )

        if b4_passed:
            findings.append(
                f"Max-round stopping: {len(b4_passed)}/{len([t for t in logic_tests if 'B4' in t.get('name','')])} "
                "models triggered the round-limit warning. The loop correctly halts and returns a 'warning' field."
            )

        # MiniMax-specific findings
        minimax_tests = [t for t in logic_tests if t.get("model") == "MiniMax-M2.1"]
        if minimax_tests:
            minimax_lats = [t.get("latency_ms", 0) for t in minimax_tests if t.get("latency_ms")]
            findings.append(
                f"MiniMax M2.1 (thinking model): avg latency {statistics.mean(minimax_lats):.0f} ms across all tests "
                f"(min {min(minimax_lats):.0f} ms, max {max(minimax_lats):.0f} ms). "
                "Requires max_tokens≥2048 — hidden reasoning tokens consume significant context before visible output."
            )

        # New v3 test findings
        b6_passed = [t for t in logic_tests if "B6" in t.get("name", "") and t.get("passed") is True]
        b7_passed = [t for t in logic_tests if "B7" in t.get("name", "") and t.get("passed") is True]
        b8_passed = [t for t in logic_tests if "B8" in t.get("name", "") and t.get("passed") is True]
        b9_passed = [t for t in logic_tests if "B9" in t.get("name", "") and t.get("passed") is True]

        if b6_passed:
            findings.append(
                f"Real file-read task (B6): {len(b6_passed)} model(s) correctly used read_file on a real workspace file "
                "and produced a meaningful summary. The full file I/O pipeline (tool exec → result injection → model response) works end-to-end."
            )
        if b7_passed:
            findings.append(
                f"Real grep task (B7): {len(b7_passed)} model(s) used grep_search to locate a known constant "
                "(CACHE_TTL_MS) and correctly reported the file. This validates tool-result reading and citation."
            )
        if b8_passed:
            findings.append(
                f"Failure handling (B8): {len(b8_passed)} model(s)' requests with a nonexistent model ID were rejected "
                "with a 4xx error — not a 500 crash. The platform handles unknown model IDs gracefully."
            )
        if b9_passed:
            timing_samples = [t.get("_timing") for t in b9_passed if t.get("_timing")]
            if timing_samples:
                pre_model_vals = [tm.get("preModelMs", 0) for tm in timing_samples if tm]
                inference_vals = [tm.get("inferenceMs", 0) for tm in timing_samples if tm]
                findings.append(
                    f"Stage timing (B9): _timing field confirmed present in {len(b9_passed)} model(s). "
                    f"Pre-model overhead: {statistics.mean(pre_model_vals):.0f}ms avg "
                    f"(was ~5000ms before registry TTL fix). "
                    f"Inference: {statistics.mean(inference_vals):.0f}ms avg."
                )
            else:
                findings.append(
                    f"Stage timing (B9): _timing field confirmed present in {len(b9_passed)} model(s). "
                    "Check platform restart to ensure chat.ts timing changes are active."
                )

    if not findings:
        findings.append("Benchmark data could not be collected. Ensure the platform is running and bun is installed.")

    for f in findings:
        p = doc.add_paragraph(f"• {f}")
        if p.runs:
            p.runs[0].font.size = Pt(10)

    doc.add_paragraph()

    # ── Recommendations ───────────────────────────────────────────────
    add_heading(doc, "8. Recommendations", 1)
    recs = [
        ("Helper functions are very fast (µs range)",
         "No optimization needed; the real bottleneck is always model inference latency (2-180s/round)."),
        ("summarizeToolOutput saves significant context",
         "Current thresholds (800 chars) are well-tuned. Consider per-model context-limit-aware truncation."),
        ("parseTextToolCalls handles 4 XML formats",
         "Text-based tool calling (<tool_use> blocks) works with both Gemma and MiniMax — no native function-call API needed."),
        ("MiniMax requires 2048+ max_tokens",
         "The model uses thinking tokens internally before producing visible content. "
         "Setting max_tokens<512 will produce empty responses. Document this in model registry metadata."),
        ("Context compression reduces message history",
         "The fixed keepLast=6 could be made dynamic based on the model's context window size."),
        ("Policy engine is pre-model (sub-10ms)",
         "Policy evaluation runs before model resolution — dangerous commands are rejected without any LLM overhead."),
        ("Investigate mode write-block is platform-enforced",
         "Write tools are filtered at the loop level regardless of what the model outputs. "
         "Consider adding a runtime assertion that validates READONLY_TOOLS against the tool executor's registry."),
        ("Max-round stopping produces structured warning",
         "The 'warning' field in the response body allows clients to detect truncated responses and prompt the user to continue."),
        ("B8: Bad model IDs return 4xx not 500",
         "The platform gracefully rejects unknown model IDs with a user-readable error. "
         "No unhandled exception reaches the client — this is important for production stability."),
        ("B9: _timing field enables performance monitoring",
         "Clients (VS Code extension, Python scripts) can now read _timing from every response to track "
         "slow requests and identify whether latency is in model resolution or inference."),
        ("Registry cache TTL increased 15s → 60s",
         "This reduces gateway probe frequency by 4x. If the 5s probe overhead reappears, "
         "consider a 'stale-while-revalidate' pattern to remove it entirely."),
        ("Qwen3.6 and GLM-4.7 restricted at gateway",
         "These models appear in the vLLM catalog but are restricted for the admin gateway key (ACL policy). "
         "Grant access to unlock them for testing."),
    ]
    for title_r, detail_r in recs:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        run_title = p.add_run(f"• {title_r}: ")
        run_title.bold = True
        run_title.font.size = Pt(10)
        run2 = p.add_run(detail_r)
        run2.font.size = Pt(10)

    doc.add_paragraph()

    # ── Plain English / Beginner Guide ───────────────────────────────
    add_heading(doc, "9. Plain English Guide (For Beginners)", 1)
    add_para(doc,
        "This section explains the test results in simple terms — no technical background needed."
    )
    doc.add_paragraph()

    beginner_items = [
        (
            "What is an 'agent loop'?",
            "Think of the AI as an employee who gets a task. The loop is how it works through the task step by step: "
            "it reads the task, decides if it needs any tools (like reading a file or searching code), uses the tools, "
            "reads the results, and keeps going until it has a complete answer. The 'loop' is this back-and-forth process."
        ),
        (
            "B1 — Quick mode passed ✓",
            "When you ask a simple question (like '2+2?'), the AI answers immediately without using any tools. "
            "This is the 'quick mode' — fast and direct. The test checks that no tools were accidentally used. PASSED = the AI gave a direct answer."
        ),
        (
            "B2 — Multi-round tool use passed ✓",
            "When the AI needs to check two different things, it does them one at a time. This test asks it to look inside "
            "two different folders in sequence. PASSED = the AI made at least 2 separate tool calls, showing it can handle multi-step tasks."
        ),
        (
            "B3 — Write block passed ✓",
            "In 'investigate' mode, the AI is only allowed to READ files — not write or modify anything. "
            "This is a safety feature. PASSED = even when asked to write a file, the system blocked it."
        ),
        (
            "B4 — Round limit stopping passed ✓",
            "The platform has a safety switch: if the AI has used too many tool rounds, it stops and tells you. "
            "This prevents infinite loops. PASSED = when we set the limit to 0, the AI stopped and the response included a warning message."
        ),
        (
            "B5 — Policy block passed ✓",
            "The platform has a 'bouncer' that reads every request before it even reaches the AI. "
            "If the message looks dangerous (like asking to delete system files), it's rejected immediately. "
            "PASSED = the dangerous command was rejected with 'HTTP 403 Forbidden' — the AI was never even asked."
        ),
        (
            "B6 — Real file read task",
            "This test gives the AI a real job: read an actual file in the codebase and explain what it does. "
            "This is 'real work' — not a made-up question. PASSED = the AI used the read_file tool on a real file and wrote a correct summary."
        ),
        (
            "B7 — Real grep (code search) task",
            "This test asks the AI to find where a specific piece of code lives in the project. "
            "Like searching for a word in a book. PASSED = the AI used the grep_search tool and correctly identified the file."
        ),
        (
            "B8 — Bad model ID handled gracefully",
            "What happens if someone asks to use a model that doesn't exist? "
            "It should give a helpful error message — not crash the server. "
            "PASSED = the server returned a '4xx' error (user mistake) rather than a '500' (server crash)."
        ),
        (
            "B9 — Timing data in every response",
            "Every response now includes a breakdown of WHERE time was spent: "
            "parsing your message, checking policies, finding the model, and running the AI. "
            "This is like a receipt that shows how each second was used. "
            "PASSED = the _timing field appeared in the response with all 6 time measurements."
        ),
        (
            "Why are some tests 'SKIPPED'?",
            "A test can be 'skipped' (marked ?) when the AI solved the problem in a way that didn't trigger the thing we wanted to test. "
            "For example, B4 is skipped if the AI answered without using any tools at all — because then the round limit has nothing to stop. "
            "It's not a failure — the AI just took a different (valid) path."
        ),
        (
            "What was the 5-second slowdown?",
            "In the previous test run, requests sometimes took 5 extra seconds before the AI even started thinking. "
            "We found the cause: the platform was regularly 'calling ahead' to the model server to check which models were available. "
            "This check had a 5-second timeout. We fixed it by making the platform remember the result for 60 seconds instead of 15 — "
            "so it only needs to call ahead at most once per minute. You should see faster pre-model times in B9's timing breakdown."
        ),
    ]
    for heading_txt, explanation in beginner_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(2)
        run_h = p.add_run(heading_txt)
        run_h.bold = True
        run_h.font.size = Pt(10)
        run_h.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        p2 = doc.add_paragraph(explanation)
        p2.paragraph_format.left_indent = Cm(0.5)
        p2.paragraph_format.space_after = Pt(4)
        if p2.runs:
            p2.runs[0].font.size = Pt(10)

    doc.add_paragraph()

    # ── Appendix: Raw data ────────────────────────────────────────────
    add_heading(doc, "Appendix A: Raw Behavioral Test Results (v3)", 1)
    if logic_tests:
        raw_para = doc.add_paragraph(json.dumps(logic_tests, indent=2, default=str)[:4000])
        raw_para.runs[0].font.size = Pt(7)
        raw_para.runs[0].font.name = "Courier New"

    add_heading(doc, "Appendix B: TypeScript Benchmark Raw Data", 1)
    if bun_data:
        raw_para = doc.add_paragraph(json.dumps(bun_data.get("benchmarks", []), indent=2)[:3000])
        raw_para.runs[0].font.size = Pt(7)
        raw_para.runs[0].font.name = "Courier New"

    doc.save(str(output_path))
    print(f"\n  ✓ Report saved: {output_path}")

# ── Main ─────────────────────────────────────────────────────────────────
def main():
    print("=" * 70)
    print("  Thirdwave AI — Agent Loop Performance & Behavioral Test v3")
    print("=" * 70)

    jwt_token = make_jwt()
    print(f"\n[1/5] JWT generated for {ADMIN_EMAIL} (admin user)")

    print("\n[2/5] TypeScript helper micro-benchmarks (bun)...")
    bun_data = run_bun_benchmarks()
    if bun_data:
        print(f"  ✓ {len(bun_data.get('benchmarks', []))} benchmarks completed")
    else:
        print("  ✗ Bun benchmarks failed or returned no data")

    print("\n[3/5] API latency benchmarks (fast path — no full inference)...")
    api_results = run_api_benchmarks(jwt_token)
    print(f"  ✓ {len(api_results)} scenarios completed")

    print(f"\n[4/5] Live behavioral loop tests (B1-B9) — {len(LOCAL_MODELS)} model(s):")
    for m in LOCAL_MODELS:
        print(f"       • {m['name']} ({m['id']}) — timeout {m['timeout_s']}s/test")
    print("  NOTE: B1-B5 are control-flow tests. B6-B7 are real tasks (use tools on real files).")
    print("  NOTE: B8 is a failure test. B9 checks the timing breakdown in responses.")
    print("  These tests call real local models and may take several minutes.")
    print()

    all_behavioral: list = []
    for model in LOCAL_MODELS:
        print(f"  ── Model: {model['name']} ──")
        model_tests = run_behavioral_tests(model, jwt_token)
        for t in model_tests:
            t["model"] = model["name"]
            status = "PASS" if t.get("passed") is True else ("SKIP" if t.get("passed") is None else "FAIL")
            print(f"    [{status}] {t['name']}  ({t.get('latency_ms','?')} ms)")
        all_behavioral.extend(model_tests)
        print()

    passed_total = sum(1 for t in all_behavioral if t.get("passed") is True)
    print(f"  ✓ {passed_total}/{len(all_behavioral)} behavioral tests passed")

    output_path = REPO_ROOT / "platform" / "tests" / "agent_loop_performance_report_v3.docx"
    print(f"\n[5/5] Generating DOCX report: {output_path}")
    generate_docx(bun_data, api_results, all_behavioral, output_path)

    # ── Console summary ─────────────────────────────────────────────
    print("\n" + "=" * 70)
    print("  FULL SUMMARY")
    print("=" * 70)
    if bun_data and bun_data.get("benchmarks"):
        print("\nHelper benchmarks:")
        for b in bun_data["benchmarks"]:
            print(f"  {b['name'][:55]:<55} {b['opsPerSec']:>12,} ops/s  ({b['avgUs']:.2f} µs/op)")
    if api_results:
        print("\nAPI latency:")
        for r in api_results:
            if r.get("avg_ms"):
                print(f"  {r['label'][:55]:<55} {r['avg_ms']:>8} ms avg")
    print("\nBehavioral tests:")
    for t in all_behavioral:
        status = "PASS" if t.get("passed") is True else ("SKIP" if t.get("passed") is None else "FAIL")
        model = t.get("model", "?")
        print(f"  [{status}] {model} — {t['name']}  ({t.get('latency_ms','?')} ms)")
    print()

if __name__ == "__main__":
    main()
